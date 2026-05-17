
# CERTIFICATE ENGINE

# This module handles the core logic for:
# 1. Loading configuration and data
# 2. Validating student records
# 3. Filling DOCX templates with student data
# 4. Converting DOCX files to PDF in batch

import json
import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
try:
    from num2words import num2words
except ImportError:
    num2words = None

# Try to import docx2pdf for conversion
try:
    from docx2pdf import convert as docx_to_pdf
except ImportError:
    docx_to_pdf = None


class CertificateGenerator:
    def __init__(self):
        """Initialize the generator by loading configs and setting up paths."""
        self.base_dir = Path(__file__).resolve().parent.parent
        self._load_config()
        self._setup_directories()
        
    def _load_config(self):
        """Load configuration files."""
        try:
            with open(self.base_dir / "config/config.json") as f:
                self.config = json.load(f)
                
            with open(self.base_dir / "config/certificate_config.json") as f:
                self.cert_config = json.load(f)
                
            self.data_file = self.base_dir / self.config["master_dataset_path"]
            self.template_file = self.base_dir / "templates" / self.cert_config["certificate_template"]
            self.placeholders = self.cert_config["placeholders"]
            self.required_fields = self.cert_config.get("required_fields", [])
            self.activity_log_file = self.base_dir / "output" / "certificate_generation_log.json"
            
            # Load official professor list for name expansion
            self.professors_file = self.base_dir / "config/professors.json"
            self.official_professors = []
            if self.professors_file.exists():
                with open(self.professors_file, 'r') as f:
                    self.official_professors = json.load(f)
            
        except Exception as e:
            print(f" Configuration error: {e}")
            raise

    def _setup_directories(self):
        """Ensure all necessary output directories exist."""
        self.pdf_folder = self.base_dir / self.config["pdf_folder"]
        self.generated_folder = self.pdf_folder / "generated"
        self.individual_folder = self.pdf_folder / "individual"
        self.docx_folder = self.pdf_folder / "docx"
        self.draft_folder = self.base_dir / self.config["draft_folder"]
        
        for folder in [self.pdf_folder, self.generated_folder, self.individual_folder, self.docx_folder, self.draft_folder]:
            folder.mkdir(parents=True, exist_ok=True)

    
    # LOGGING
    

    def log_activity(self, activity_type, data):
        """Append an entry to the activity log JSON file."""
        self.activity_log_file.parent.mkdir(parents=True, exist_ok=True)

        entry = {
            "timestamp": datetime.now().isoformat(),
            "type": activity_type,
            "data": data
        }

        try:
            logs = []
            if self.activity_log_file.exists():
                with open(self.activity_log_file, 'r') as f:
                    logs = json.load(f)
            
            logs.append(entry)

            with open(self.activity_log_file, 'w') as f:
                json.dump(logs, f, indent=2)

        except Exception as e:
            print(f"Error logging activity: {e}")

    
    # DATA VALIDATION & FORMATTING
    

    def is_complete(self, row):
        """Check if a student row has all required fields."""
        for field in self.required_fields:
            if field not in row or not row[field] or pd.isna(row[field]):
                return False
        return True

    def _num_to_words(self, n):
        """Converts an integer to words (e.g. 6 -> six) using num2words library."""
        try:
            val = int(float(n))
            if num2words:
                return num2words(val)
            print(f"WARNING: num2words not installed. Returning digit for {val}.")
            return str(val)
        except Exception as e:
            print(f"ERROR: input '{n}' failed num2words conversion: {e}")
            return str(val)

    def _get_value(self, column, row):
        """Get formatted value for a placeholder column."""
        try:
            val = row.get(column, "") if hasattr(row, 'get') else ""

            if pd.isna(val):
                return ""

            # Format Numbers to Words
            if column in self.cert_config.get("convert_to_words", []):
                return self._num_to_words(val)

            # Format Dates
            date_fmt = self.cert_config.get("date_format", "%d/%m/%Y")
            
            if isinstance(val, (pd.Timestamp, datetime)):
                return val.strftime(date_fmt)
                
            if "date" in column.lower() and isinstance(val, str):
                try:
                    # Parse string (e.g., '2023-05-22') and format it
                    parsed_date = pd.to_datetime(val)
                    return parsed_date.strftime(date_fmt)
                except Exception:
                    pass

            result_str = str(val)
            
            if column in self.cert_config.get("force_uppercase", []):
                return result_str.upper()
                
            if column == "supervisor_name" and result_str:
                # Try to expand shortened names using official list
                for full_name in self.official_professors:
                    if result_str.lower() in full_name.lower():
                        return full_name
            
            if column in ["focused_on", "contributed_towards"] and result_str:
                stop_words = {"and", "or", "the", "a", "an", "in", "on", "at", "to", "for", "of", "with", "by"}
                words = result_str.split()
                title_words = []
                for i, word in enumerate(words):
                    if i > 0 and word.lower() in stop_words:
                        title_words.append(word.lower())
                    else:
                        title_words.append(word.capitalize())
                return " ".join(title_words)

            return result_str

        except Exception:
            return ""

    
    # DOCX TEMPLATE ENGINE
    

    def _fix_broken_placeholders(self, paragraph):
        """
        Merges adjacent text runs to ensure placeholders like {{Name}} 
        are not split across multiple runs (e.g., {{ + Name + }}).
        """
        while True:
            text = paragraph.text
            if not ("{{" in text and "}}" in text):
                break
                
            runs = paragraph.runs
            start_idx, end_idx = -1, -1
            
            # Find start of placeholder "{{"
            for i, run in enumerate(runs):
                if "{{" in run.text and "}}" not in run.text:
                    start_idx = i
                    break
            
            if start_idx != -1:
                # Find end of placeholder "}}"
                for j in range(start_idx + 1, len(runs)):
                    if "}}" in runs[j].text:
                        end_idx = j
                        break
                
                if end_idx != -1:
                    # Check for bold formatting in any part of the split tag
                    is_bold = any(runs[k].font.bold for k in range(start_idx, end_idx + 1))

                    # Merge text into the first run
                    for k in range(start_idx + 1, end_idx + 1):
                        runs[start_idx].text += runs[k].text
                        runs[k].text = "" # Clear merged runs
                    
                    # Restore bold if applicable
                    if is_bold:
                        runs[start_idx].font.bold = True
                    
                    continue # Check for next placeholder
            
            break # No more split placeholders found

    def _replace_in_paragraph(self, paragraph, replacements, bold_tnr_keys=None, bold_calibri_keys=None):
        """
        Apply replacements run-by-run to perfectly preserve the original styling 
        of static text (like quotation marks) that might be in different fonts.
        """
        import re
        
        if not any(key in paragraph.text for key in replacements):
            return

        self._fix_broken_placeholders(paragraph)
        
        sorted_keys = sorted(replacements.keys(), key=len, reverse=True)
        pattern = "|".join(re.escape(k) for k in sorted_keys)
        
        new_runs_data = []
        
        for run in paragraph.runs:
            if not any(key in run.text for key in replacements):
                new_runs_data.append({
                    'text': run.text,
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'name': run.font.name,
                    'size': run.font.size
                })
                continue
                
            last_idx = 0
            for match in re.finditer(pattern, run.text):
                start, end = match.span()
                
                # Add preceding static text with original run's formatting
                if start > last_idx:
                    new_runs_data.append({
                        'text': run.text[last_idx:start],
                        'bold': run.font.bold,
                        'italic': run.font.italic,
                        'name': run.font.name,
                        'size': run.font.size
                    })
                
                key = match.group()
                val = str(replacements[key])
                is_tnr_bold = bold_tnr_keys and key in bold_tnr_keys
                is_calibri_bold = bold_calibri_keys and key in bold_calibri_keys
                is_bold = is_tnr_bold or is_calibri_bold
                
                # Period grouping logic (if period is inside the same run)
                next_chars = run.text[end:end+2]
                if is_bold and next_chars.startswith('.'):
                    val += "."
                    last_idx = end + 1
                else:
                    last_idx = end
                    
                # Add the replacement text with special formatting
                new_runs_data.append({
                    'text': val,
                    'bold': True if is_bold else run.font.bold,
                    'italic': run.font.italic,
                    'name': 'Times New Roman' if is_tnr_bold else ('Calibri' if is_calibri_bold else run.font.name),
                    'size': Pt(14) if is_bold else run.font.size
                })
                
            # Add any remaining static text in the run
            if last_idx < len(run.text):
                new_runs_data.append({
                    'text': run.text[last_idx:],
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'name': run.font.name,
                    'size': run.font.size
                })

        # 3. Clear existing runs and rebuild
        paragraph.text = ""
        for r_data in new_runs_data:
            if not r_data['text']:
                continue
            run = paragraph.add_run(r_data['text'])
            if r_data['bold'] is not None: run.font.bold = r_data['bold']
            if r_data['italic'] is not None: run.font.italic = r_data['italic']
            if r_data['name'] is not None: run.font.name = r_data['name']
            if r_data['size'] is not None: run.font.size = r_data['size']

    def create_docx(self, row, student_id):
        """Generate a single DOCX certificate for a student."""
        try:
            if not self.template_file.exists():
                raise FileNotFoundError(f"Template not found: {self.template_file}")

            doc = Document(self.template_file)
            
            # Prepare data mapping
            replacements = {
                tag: self._get_value(col, row)
                for tag, col in self.placeholders.items()
            }
            
            # Automatically fill issue date with the current date in specific format (DD. MM. YYYY)
            today_date_issue = datetime.now().strftime("%d. %m. %Y")
            replacements["{{issue_date}}"] = today_date_issue
            replacements["{{Issue_Date}}"] = today_date_issue
            replacements["{{Issue_date}}"] = today_date_issue
            replacements["{{ISSUE_DATE}}"] = today_date_issue
            
            # Prepare bold keys based on config
            bold_tnr_keys = set()
            bold_calibri_keys = set()
            for tag, col in self.placeholders.items():
                is_forced_upper = col in self.cert_config.get("force_uppercase", [])
                is_forced_bold = col in self.cert_config.get("force_bold", [])
                
                if col in ["focused_on", "contributed_towards"]:
                    bold_calibri_keys.add(tag)  # Set to Calibri instead of TNR
                else:
                    if is_forced_bold:
                        bold_tnr_keys.add(tag)
                    if is_forced_upper:
                        bold_calibri_keys.add(tag)
            
            # Apply to body
            for p in doc.paragraphs:
                self._replace_in_paragraph(p, replacements, bold_tnr_keys, bold_calibri_keys)

            # Apply to tables
            for table in doc.tables:
                for row_obj in table.rows:
                    for cell in row_obj.cells:
                        for p in cell.paragraphs:
                            self._replace_in_paragraph(p, replacements, bold_tnr_keys, bold_calibri_keys)
            
            output_path = self.docx_folder / f"{student_id}_certificate.docx"
            doc.save(output_path)
            return output_path

        except Exception as e:
            print(f"Error creating DOCX for {student_id}: {e}")
            self.log_activity("certificate_failed", {"student_id": str(student_id), "error": str(e)})
            return None

    
    # PDF CONVERSION
    

    def batch_convert_pdf(self, source_dir, output_dir):
        """Convert all DOCX files in a folder to PDF using win32com for maximum image quality."""
        try:
            import win32com.client
            import pythoncom
            from pathlib import Path
            
            # Initialize COM in the current thread
            pythoncom.CoInitialize()
            
            print(f"Converting files in {source_dir} to PDF (High Quality Mode)...")
            
            # Use DispatchEx to get a clean Word instance
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            
            wdFormatPDF = 17
            wdExportOptimizeForPrint = 0
            
            source_path = Path(source_dir).resolve()
            output_path = Path(output_dir).resolve()
            
            success = True
            for docx_file in source_path.glob("*.docx"):
                pdf_file = output_path / (docx_file.stem + ".pdf")
                try:
                    doc = word.Documents.Open(str(docx_file), ReadOnly=True)
                    # ExportAsFixedFormat explicitly forces high-quality print optimization
                    doc.ExportAsFixedFormat(
                        OutputFileName=str(pdf_file),
                        ExportFormat=wdFormatPDF,
                        OpenAfterExport=False,
                        OptimizeFor=wdExportOptimizeForPrint,
                        Range=0,  # wdExportAllDocument
                        Item=0,   # wdExportDocumentContent
                        IncludeDocProps=True,
                        KeepIRM=True,
                        CreateBookmarks=0,  # wdExportCreateNoBookmarks
                        DocStructureTags=True,
                        BitmapMissingFonts=True,
                        UseISO19005_1=False
                    )
                    doc.Close(False)
                except Exception as e:
                    print(f"Error converting {docx_file.name}: {e}")
                    success = False
                    
            word.Quit()
            pythoncom.CoUninitialize()
            return success, None
            
        except Exception as e:
            print(f"Batch conversion failed: {e}")
            return False, str(e)

    
    # MAIN WORKFLOW
    

    def run_batch(self, student_ids=None):
        """Execute the full batch generation process."""
        try:
            print("=" * 60)
            if student_ids:
                print(f"BATCH CERTIFICATE GENERATION STARTED FOR IDs: {', '.join(student_ids)}")
            else:
                print("BATCH CERTIFICATE GENERATION STARTED FOR ALL")
            print("=" * 60)

            self.log_activity("generation_started", {"student_ids": student_ids})
            
            # Use Database Manager instead of Excel
            from db_manager import DatabaseManager
            db = DatabaseManager()
            conn = db.get_connection()
            cursor = conn.cursor(dictionary=True)

            # Prepare query
            if student_ids:
                placeholders = ','.join(['%s'] * len(student_ids))
                query = f"SELECT * FROM students WHERE student_id IN ({placeholders})"
                cursor.execute(query, [str(sid) for sid in student_ids])
                rows = cursor.fetchall()
            else:
                query = "SELECT * FROM students WHERE is_complete = 1"
                cursor.execute(query)
                rows = cursor.fetchall()
            
            students_to_process = []
            for row in rows:
                data = dict(row)
                if self.is_complete(data):
                    students_to_process.append((data["student_id"], data))

            if not students_to_process:
                 conn.close()
                 return {"status": "success", "message": "No students found matching the criteria (Ensure IDs are correct and fields complete)"}

            # STEP 1: Generate DOCX
            docx_count = 0
            failed_count = 0
            
            # Clean old DOCX files
            for f in self.docx_folder.glob("*.docx"):
                try: f.unlink()
                except: pass

            for student_id, row in students_to_process:
                if self.create_docx(row, student_id):
                    docx_count += 1
                    # Update DB with generation status
                    try:
                        ts = datetime.now().isoformat()
                        cursor.execute("UPDATE students SET certificate_generated = 1, generated_at = %s WHERE student_id = %s", (ts, str(student_id)))
                        conn.commit()
                    except Exception as e:
                        print(f"Error updating generation status: {e}")
                    self.log_activity("docx_generated", {"student_id": str(student_id)})
                else:
                    failed_count += 1
            
            conn.close() # Close DB connection

            # STEP 2: Convert to PDF
            # Determine output folder: 'individual' if specific IDs provided, else 'generated'
            output_folder = self.individual_folder if student_ids else self.generated_folder
            
            if docx_count > 0:
                self.log_activity("conversion_started", {"count": docx_count, "target": output_folder.name})
                success, error = self.batch_convert_pdf(self.docx_folder, output_folder)
                
                if success:
                    self.log_activity("batch_conversion_complete", {"count": docx_count})
                else:
                    self.log_activity("batch_conversion_failed", {"error": error})

            print("=" * 60)
            print(f"Generated: {docx_count}")
            print(f"Failed:    {failed_count}")
            print("=" * 60)

            return {
                "status": "success",
                "generated": docx_count,
                "failed": failed_count,
                "output_folder": output_folder.name
            }

        except Exception as e:
            print(f"Critical Error: {e}")
            return {"status": "error", "error": str(e)}



# BACKWARD COMPATIBILITY WRAPPER


def generate_all_certificates(student_ids=None):
    """Wrapper function to maintain compatibility with existing API calls."""
    engine = CertificateGenerator()
    return engine.run_batch(student_ids)


# MAIN ENTRY POINT


if __name__ == "__main__":
    import sys
    ids = sys.argv[1:] if len(sys.argv) > 1 else None
    result = generate_all_certificates(ids)
    print(json.dumps(result, indent=2))
