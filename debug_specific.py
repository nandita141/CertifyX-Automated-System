import pandas as pd
from backend.certificate_engine import CertificateGenerator
import sys

def debug_student():
    cg = CertificateGenerator()
    df = pd.read_excel(cg.data_file)
    
    # search for student
    target_name = "P V Ramachandran" # from screenshot
    # Or search for any student with no_of_weeks == 6
    
    # Filter
    # Try exact match first
    student = df[df['student_name'].str.contains("Ramachandran", case=False, na=False)]
    
    if student.empty:
        print("Student not found by name. searching for weeks=6")
        student = df[df['no_of_weeks'] == 6]
        
    if student.empty:
        print("No student found with 6 weeks.")
        return

    row = student.iloc[0]
    s_id = row['student_id']
    print(f"Found Student ID: {s_id}")
    print(f"Name: {row['student_name']}")
    print(f"Weeks in DB: {row['no_of_weeks']} (Type: {type(row['no_of_weeks'])})")
    
    # Check conversion logic manually
    val = cg._get_value('no_of_weeks', row)
    print(f"Converted Value from engine: '{val}'")
    
    # Check config
    print(f"Config convert_to_words: {cg.cert_config.get('convert_to_words')}")
    
    # Check template placeholders
    try:
        from docx import Document
        doc = Document(cg.template_file)
        placeholders = set()
        for p in doc.paragraphs:
            if "{{" in p.text:
                placeholders.add(p.text)
            for r in p.runs:
                 if "{{" in r.text:
                     placeholders.add(r.text)
        print(f"Template Placeholders (snippet): {list(placeholders)[:5]}")
        
        # Check specifically for NO_OF_WEEKS
        found_tag = False
        tag_variations = ["{{NO_OF_WEEKS}}", "{{no_of_weeks}}", "{{No_Of_Weeks}}"]
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        for t in tag_variations:
            if t in full_text:
                print(f"FOUND TAG IN TEMPLATE: {t}")
                found_tag = True
        
        if not found_tag:
            print("WARNING: {{NO_OF_WEEKS}} NOT FOUND IN TEMPLATE TEXT!")
            # Print text around "weeks"
            import re
            matches = re.findall(r"(.{20}weeks.{20})", full_text)
            print("Context around 'weeks' in template:")
            for m in matches:
                print(f"  ...{m}...")
                
    except Exception as e:
        print(f"Template inspection failed: {e}")

if __name__ == "__main__":
    debug_student()
